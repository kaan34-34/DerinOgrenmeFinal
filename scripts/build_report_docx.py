from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Beyin_Tumoru_YOLOv8_Proje_Raporu.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_cell_margins(table)


def style_run(run, bold=False, italic=False, color=None, size=None, font="Calibri"):
    run.bold = bold
    run.italic = italic
    run.font.name = font
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    color = "2E74B5" if level <= 2 else "1F4D78"
    size = 16 if level == 1 else 13 if level == 2 else 12
    for run in p.runs:
        style_run(run, bold=True, color=color, size=size)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12 if level == 2 else 8)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6 if level == 2 else 4)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    for run in p.runs:
        style_run(run, size=11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        style_run(run, size=11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        style_run(run, size=11)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F4F7")
    p._p.get_or_add_pPr().append(shading)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "F2F4F7")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                style_run(run, bold=True, size=10)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    style_run(run, size=10)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_figure(doc, image_path, caption):
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        style_run(run, italic=True, color="555555", size=9)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Beyin Tümörü Tespiti için YOLOv8 Tabanlı Derin Öğrenme Projesi")
    style_run(r, bold=True, color="0B2545", size=20)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Yüksek Lisans Derin Öğrenme Proje Raporu")
    style_run(r, color="555555", size=12)

    meta = [
        ("Ders", "Derin Öğrenme"),
        ("Proje Türü", "Nesne tespiti, medikal görüntü analizi"),
        ("Model", "YOLOv8n"),
        ("Veri Seti", "Brain Tumor YOLO Format Dataset"),
        ("Ortam", "Python, Ultralytics YOLOv8, PyTorch, macOS CPU"),
    ]
    add_table(doc, ["Alan", "Değer"], meta, widths=[1.6, 4.8])

    add_heading(doc, "Özet", 1)
    add_body(
        doc,
        "Bu projede beyin görüntülerinde tümörle ilişkili bölgelerin otomatik olarak "
        "tespit edilmesi amaçlanmıştır. Problem, görüntü sınıflandırma yerine nesne "
        "tespiti olarak ele alınmış; modelden hem sınıf hem de sınırlayıcı kutu "
        "tahmini üretmesi beklenmiştir.",
    )
    add_body(
        doc,
        "Çalışmada Ultralytics YOLOv8 ailesinin hafif sürümü olan YOLOv8n kullanılmıştır. "
        "Veri seti doğrulanmış, eksik etiket dosyaları tamamlanmış ve eğitim, değerlendirme "
        "ve tahmin süreçlerini çalıştıran Python scriptleri hazırlanmıştır.",
    )

    add_heading(doc, "1. Projenin Amacı ve Hedefleri", 1)
    add_body(
        doc,
        "Projenin amacı, YOLOv8 tabanlı bir nesne tespit modeli ile beyin görüntülerindeki "
        "pozitif ve negatif bölgeleri otomatik olarak belirlemektir. Bu amaç medikal "
        "görüntü analizinde derin öğrenme modellerinin uçtan uca uygulanmasını göstermektedir.",
    )
    for item in [
        "YOLO formatındaki veri seti yapısını incelemek.",
        "Görüntü ve etiket eşleşmelerini doğrulamak.",
        "Eksik etiket dosyalarını uygun biçimde tamamlamak.",
        "YOLOv8n için eğitim yapılandırmasını hazırlamak.",
        "Eğitim, değerlendirme ve tahmin adımlarını scriptlerle otomatikleştirmek.",
        "Elde edilen çıktıları ve metrikleri raporlamak.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. Veri Seti", 1)
    add_body(
        doc,
        "Veri seti proje klasöründe brain-tumor adıyla bulunmaktadır ve YOLO nesne "
        "tespiti formatına uygun olarak images/train, images/val, labels/train ve "
        "labels/val klasörlerinden oluşmaktadır.",
    )
    add_table(
        doc,
        ["Bölüm", "Görüntü", "Etiket", "Kutu", "Boş Etiket"],
        [
            ["Eğitim", 893, 893, 925, 15],
            ["Doğrulama", 223, 223, 241, 0],
            ["Toplam", 1116, 1116, 1166, 15],
        ],
        widths=[1.5, 1.1, 1.1, 1.1, 1.4],
    )
    add_table(
        doc,
        ["Bölüm", "Negative Kutu", "Positive Kutu"],
        [["Eğitim", 437, 488], ["Doğrulama", 154, 87]],
        widths=[2.2, 2.0, 2.0],
    )
    add_body(
        doc,
        "İlk kontrolde eğitim setinde 15 görüntünün etiket dosyası eksikti. YOLO formatında "
        "boş etiket dosyası, görüntüde işaretlenmiş nesne bulunmadığını gösterdiği için bu "
        "dosyalar boş .txt olarak oluşturulmuş ve veri bütünlüğü sağlanmıştır.",
    )

    add_heading(doc, "3. Yöntem", 1)
    add_heading(doc, "3.1 Model Mimarisi", 2)
    add_body(
        doc,
        "YOLOv8, tek aşamalı nesne tespit yaklaşımına dayanan modern bir mimaridir. Model "
        "görüntü girdisinden sınıf, güven skoru ve sınırlayıcı kutu koordinatları üretir. "
        "Bu projede hızlı prototipleme ve CPU üzerinde çalışabilirlik nedeniyle YOLOv8n "
        "tercih edilmiştir.",
    )
    add_heading(doc, "3.2 Etiket Formatı", 2)
    add_body(doc, "YOLO etiket formatı normalize edilmiş beş sütundan oluşur:")
    add_code(doc, "class_id x_center y_center width height")

    add_heading(doc, "4. Uygulama ve Proje Dosyaları", 1)
    add_table(
        doc,
        ["Dosya", "Açıklama"],
        [
            ["brain-tumor.yaml", "YOLOv8 veri seti yapılandırması"],
            ["requirements.txt", "Proje bağımlılıkları"],
            ["validate_dataset.py", "Görüntü, etiket ve koordinat doğrulaması"],
            ["train.py", "YOLOv8n eğitim scripti"],
            ["evaluate.py", "Eğitilmiş modelin değerlendirilmesi"],
            ["predict.py", "Yeni görüntüler üzerinde tahmin yapılması"],
            ["README.md", "Çalıştırma yönergeleri"],
        ],
        widths=[2.1, 4.1],
    )

    add_heading(doc, "5. Çalıştırma Adımları", 1)
    add_number(doc, "Sanal ortam ve bağımlılık kurulumu")
    add_code(doc, "python3 -m venv .venv\n.venv/bin/python -m pip install -r requirements.txt")
    add_number(doc, "Veri seti doğrulama")
    add_code(doc, ".venv/bin/python validate_dataset.py")
    add_number(doc, "50 epoch eğitim")
    add_code(doc, ".venv/bin/python train.py --model yolov8n.pt --epochs 50 --imgsz 640 --batch 8 --device cpu")
    add_number(doc, "Değerlendirme")
    add_code(doc, ".venv/bin/python evaluate.py --weights runs/detect/brain_tumor_yolov8n/weights/best.pt")
    add_number(doc, "Tahmin")
    add_code(doc, ".venv/bin/python predict.py --weights runs/detect/brain_tumor_yolov8n/weights/best.pt --source brain-tumor/images/val")

    add_heading(doc, "6. Deneysel Sonuçlar", 1)
    add_body(
        doc,
        "Ana deney 50 epoch, 640x640 görüntü boyutu ve batch=8 ayarıyla tamamlanmıştır. "
        "Çıktılar runs/detect/brain_tumor_yolov8n klasöründe oluşmuş; en iyi ağırlık "
        "weights/best.pt, son epoch ağırlığı ise weights/last.pt olarak kaydedilmiştir.",
    )
    add_table(
        doc,
        ["Metrik", "Değer"],
        [
            ["Toplam Epoch", "50"],
            ["Son Epoch", "50"],
            ["Son Epoch Train Box Loss", "0.72275"],
            ["Son Epoch Train Class Loss", "0.78114"],
            ["Son Epoch Train DFL Loss", "0.92592"],
            ["Son Epoch Precision", "0.42183"],
            ["Son Epoch Recall", "0.83326"],
            ["Son Epoch mAP50", "0.46390"],
            ["Son Epoch mAP50-95", "0.32973"],
            ["Son Epoch Val Box Loss", "0.94503"],
            ["Son Epoch Val Class Loss", "1.33419"],
            ["Son Epoch Val DFL Loss", "1.02164"],
        ],
        widths=[3.2, 2.8],
    )
    add_body(
        doc,
        "results.csv üzerinde yapılan incelemede en yüksek mAP50-95 değerinin 37. epochta "
        "elde edildiği görülmüştür. Bu epoch, best.pt ağırlığının temsil ettiği en iyi "
        "doğrulama performansını vermektedir.",
    )
    add_table(
        doc,
        ["Metrik", "Değer"],
        [
            ["En İyi Epoch", "37"],
            ["Precision", "0.44836"],
            ["Recall", "0.86211"],
            ["mAP50", "0.49092"],
            ["mAP50-95", "0.36469"],
            ["Val Box Loss", "0.94280"],
            ["Val Class Loss", "1.08693"],
            ["Val DFL Loss", "1.01277"],
        ],
        widths=[3.2, 2.8],
    )
    add_body(
        doc,
        "best.pt ile yapılan ek doğrulamada sınıf bazlı performans aşağıdaki gibi ölçülmüştür.",
    )
    add_table(
        doc,
        ["Sınıf", "Görüntü", "Instance", "Precision", "Recall", "mAP50", "mAP50-95"],
        [
            ["Tüm sınıflar", "223", "241", "0.455", "0.847", "0.491", "0.365"],
            ["Negative", "142", "154", "0.578", "0.844", "0.601", "0.448"],
            ["Positive", "81", "87", "0.333", "0.849", "0.381", "0.281"],
        ],
        widths=[1.35, 0.85, 0.9, 0.95, 0.8, 0.8, 0.95],
    )
    add_body(
        doc,
        "Pozitif sınıf recall değerinin 0.849 olması, modelin pozitif örneklerin büyük kısmını "
        "yakalayabildiğini göstermektedir. Pozitif sınıf precision değerinin 0.333 olması ise "
        "yanlış pozitif tahminleri azaltmak için ek iyileştirme ihtiyacına işaret etmektedir.",
    )

    add_figure(
        doc,
        ROOT / "runs/detect/brain_tumor_yolov8n/results.png",
        "Şekil 1. 50 epoch eğitim ve doğrulama eğrileri.",
    )
    add_figure(
        doc,
        ROOT / "runs/detect/brain_tumor_yolov8n/confusion_matrix.png",
        "Şekil 2. 50 epoch sonrası karmaşıklık matrisi.",
    )
    add_figure(
        doc,
        ROOT / "runs/detect/predict-2/val_1 (144).jpg",
        "Şekil 3. Eğitilmiş ağırlıkla üretilen örnek tahmin çıktısı.",
    )

    add_heading(doc, "7. Tartışma", 1)
    add_body(
        doc,
        "Proje, YOLOv8 tabanlı nesne tespiti hattının medikal görüntü verisine uygulanabilir "
        "olduğunu göstermektedir. Veri setindeki eksik etiketlerin tespit edilip tamamlanması, "
        "eğitim öncesi veri doğrulamanın önemini ortaya koymuştur.",
    )
    add_body(
        doc,
        "50 epoch eğitim sonunda model genel mAP50 değerini yaklaşık 0.491, mAP50-95 değerini "
        "yaklaşık 0.365 seviyesine taşımıştır. En iyi mAP50-95 değerinin 37. epochta oluşması, "
        "sonraki epochlarda doğrulama performansında dalgalanma olduğunu göstermektedir.",
    )
    add_body(
        doc,
        "Pozitif sınıf recall değerinin yüksek olması medikal bağlamda olumludur; model pozitif "
        "örneklerin büyük bölümünü yakalayabilmiştir. Buna karşın pozitif sınıf precision değerinin "
        "düşük kalması, yanlış pozitiflerin azaltılması için eşik ayarı, veri artırma veya daha "
        "güçlü model varyantlarıyla ek deney yapılması gerektiğini göstermektedir.",
    )

    add_heading(doc, "8. Sonuç ve Gelecek Çalışmalar", 1)
    add_body(
        doc,
        "Bu çalışma kapsamında beyin tümörü tespiti için YOLOv8 tabanlı bir derin öğrenme projesi "
        "hazırlanmış, veri seti doğrulanmış ve eğitim, değerlendirme ve tahmin süreçleri çalışır "
        "hale getirilmiştir. Nihai 50 epoch eğitim tamamlanmış ve best.pt ağırlığı doğrulama seti "
        "üzerinde değerlendirilmiştir.",
    )
    for item in [
        "YOLOv8s veya YOLOv8m gibi daha büyük modellerle karşılaştırma yapılması.",
        "Pozitif sınıf için güven eşiği analiziyle yanlış pozitif ve yanlış negatif dengesinin optimize edilmesi.",
        "Veri artırma stratejilerinin precision ve recall üzerindeki etkisinin incelenmesi.",
        "Daha sağlam değerlendirme için çapraz doğrulama yaklaşımlarının denenmesi.",
        "Pozitif sınıf için ek veri veya daha dengeli örnekleme stratejilerinin denenmesi.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Kaynaklar ve Kullanılan Teknolojiler", 1)
    for item in ["Ultralytics YOLOv8", "PyTorch", "OpenCV", "Matplotlib", "NumPy", "YOLO nesne tespiti etiket formatı"]:
        add_bullet(doc, item)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Beyin Tümörü YOLOv8 Proje Raporu")
    style_run(run, color="777777", size=9)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
